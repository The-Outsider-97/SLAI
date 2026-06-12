"""Privacy-safe local document extractor for DocuMaster.

The SLAI Reader Agent is preferred by DocumentAIService.  This extractor is the
local degraded path and utility layer for word counts/conversion. It validates
file type, file size, page count, and readable text without storing content.
"""

from __future__ import annotations

import io
import os
import re
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, BinaryIO, Dict, List

from bs4 import BeautifulSoup
from docx import Document
from pypdf import PdfReader
from pypdf.errors import PdfReadError
from werkzeug.datastructures import FileStorage
from werkzeug.utils import secure_filename

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".html", ".htm", ".xml", ".odt"}
DEFAULT_MAX_FILE_SIZE_BYTES = int(os.getenv("DOCMASTER_MAX_UPLOAD_BYTES", str(20 * 1024 * 1024)))
DEFAULT_MAX_TEXT_CHARS = int(os.getenv("DOCMASTER_MAX_TEXT_CHARS", "120000"))
DEFAULT_MAX_PAGES = int(os.getenv("DOCMASTER_MAX_PAGES", "250"))


class DocumentExtractionError(ValueError):
    """Raised when a document cannot be safely read or validated."""


@dataclass(frozen=True)
class ExtractedDocument:
    filename: str
    detected_type: str
    text: str
    metadata: Dict[str, Any] = field(default_factory=dict)
    warnings: List[str] = field(default_factory=list)


class DocumentExtractor:
    """Extract readable text from supported document formats."""

    def __init__(
        self,
        *,
        max_file_size_bytes: int = DEFAULT_MAX_FILE_SIZE_BYTES,
        max_text_chars: int = DEFAULT_MAX_TEXT_CHARS,
        max_pages: int = DEFAULT_MAX_PAGES,
        allowed_extensions: set[str] | None = None,
    ) -> None:
        self.max_file_size_bytes = int(max_file_size_bytes)
        self.max_text_chars = int(max_text_chars)
        self.max_pages = int(max_pages)
        self.allowed_extensions = {ext.lower() for ext in (allowed_extensions or SUPPORTED_EXTENSIONS)}

    def extract_upload(self, file: FileStorage) -> ExtractedDocument:
        filename = secure_filename(file.filename or "uploaded_document")
        raw = file.read()
        return self.extract_bytes(raw, filename=filename)

    def extract_path(self, path: str | Path) -> ExtractedDocument:
        file_path = Path(path)
        if not file_path.exists() or not file_path.is_file():
            raise DocumentExtractionError("Selected file does not exist or is not a file.")
        if file_path.stat().st_size > self.max_file_size_bytes:
            limit_mb = self.max_file_size_bytes / (1024 * 1024)
            raise DocumentExtractionError(f"File is too large. Maximum upload size is {limit_mb:.0f} MB.")
        return self.extract_bytes(file_path.read_bytes(), filename=file_path.name)

    def extract_bytes(self, raw: bytes, *, filename: str) -> ExtractedDocument:
        safe_name = secure_filename(filename or "document")
        ext = Path(safe_name).suffix.lower()
        warnings: List[str] = []

        if ext not in self.allowed_extensions:
            raise DocumentExtractionError(
                f"Unsupported file type '{ext or 'unknown'}'. Supported: {', '.join(sorted(self.allowed_extensions))}."
            )
        if len(raw) > self.max_file_size_bytes:
            limit_mb = self.max_file_size_bytes / (1024 * 1024)
            raise DocumentExtractionError(f"File is too large. Maximum upload size is {limit_mb:.0f} MB.")
        if not raw:
            raise DocumentExtractionError("The uploaded file is empty.")

        stream = io.BytesIO(raw)
        metadata: Dict[str, Any] = {"file_size_bytes": len(raw)}
        try:
            if ext == ".pdf":
                text, page_count, page_texts, pdf_warnings = self._extract_pdf(stream)
                warnings.extend(pdf_warnings)
                metadata.update({"page_count": page_count, "page_texts": page_texts})
            elif ext == ".docx":
                text, paragraph_count, table_count = self._extract_docx(stream)
                metadata.update({"paragraph_count": paragraph_count, "table_count": table_count})
            elif ext in {".html", ".htm", ".xml"}:
                text = self._extract_markup(raw)
            elif ext == ".odt":
                text = self._extract_odt(raw)
            else:
                text = self._decode_text(raw)
        except DocumentExtractionError:
            raise
        except PdfReadError as exc:
            raise DocumentExtractionError("PDF could not be read. It may be encrypted, corrupt, or image-only.") from exc
        except Exception as exc:  # noqa: BLE001
            raise DocumentExtractionError(f"Document could not be processed: {type(exc).__name__}.") from exc

        text = self._clean_text(text)
        if not text:
            raise DocumentExtractionError("No readable text was found. The file may be scanned, encrypted, or unsupported.")
        if len(text) > self.max_text_chars:
            text = text[: self.max_text_chars]
            warnings.append(f"Text was truncated to {self.max_text_chars:,} characters for safe processing.")

        metadata.update(
            {
                "word_count": len(text.split()),
                "char_count": len(text),
                "detected_type": ext.lstrip("."),
                "language": self._guess_language(text),
                "reader_agent_used": False,
            }
        )
        return ExtractedDocument(safe_name, ext.lstrip("."), text, metadata, list(dict.fromkeys(warnings)))

    def _extract_pdf(self, stream: BinaryIO) -> tuple[str, int, List[Dict[str, Any]], List[str]]:
        reader = PdfReader(stream)
        if getattr(reader, "is_encrypted", False):
            raise DocumentExtractionError("Encrypted PDFs are not processed by default.")
        page_count = len(reader.pages)
        if page_count > self.max_pages:
            raise DocumentExtractionError(f"PDF has {page_count} pages. The configured limit is {self.max_pages} pages.")
        chunks: List[str] = []
        page_texts: List[Dict[str, Any]] = []
        warnings: List[str] = []
        for index, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text() or ""
            except Exception:  # noqa: BLE001
                page_text = ""
                warnings.append(f"Page {index + 1} could not be extracted.")
            cleaned = self._clean_text(page_text)
            page_texts.append({"page_number": index + 1, "char_count": len(cleaned), "preview": cleaned[:240]})
            if cleaned:
                chunks.append(cleaned)
        return "\n\n".join(chunks), page_count, page_texts, warnings

    @staticmethod
    def _extract_docx(stream: BinaryIO) -> tuple[str, int, int]:
        doc = Document(stream)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        table_cells: List[str] = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text and cell.text.strip():
                        table_cells.append(cell.text.strip())
        return "\n".join(paragraphs + table_cells), len(paragraphs), len(doc.tables)

    @staticmethod
    def _extract_odt(raw: bytes) -> str:
        try:
            with zipfile.ZipFile(io.BytesIO(raw)) as archive:
                content = archive.read("content.xml")
        except Exception as exc:  # noqa: BLE001
            raise DocumentExtractionError("ODT file could not be opened.") from exc
        soup = BeautifulSoup(content, "xml")
        return soup.get_text(" ")

    @staticmethod
    def _extract_markup(raw: bytes) -> str:
        decoded = DocumentExtractor._decode_text(raw)
        soup = BeautifulSoup(decoded, "html.parser")
        for tag in soup(["script", "style"]):
            tag.decompose()
        return soup.get_text(" ")

    @staticmethod
    def _decode_text(raw: bytes) -> str:
        for encoding in ("utf-8", "utf-16", "latin-1"):
            try:
                return raw.decode(encoding)
            except UnicodeDecodeError:
                continue
        return raw.decode("utf-8", errors="ignore")

    @staticmethod
    def _clean_text(text: str) -> str:
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    @staticmethod
    def _guess_language(text: str) -> str:
        sample = f" {text[:3000].lower()} "
        nl = sum(marker in sample for marker in (" de ", " het ", " een ", " voor ", " zijn ", " worden ", " niet "))
        en = sum(marker in sample for marker in (" the ", " and ", " for ", " with ", " this ", " that ", " not "))
        if nl > en:
            return "nl"
        if en > nl:
            return "en"
        return "unknown"
